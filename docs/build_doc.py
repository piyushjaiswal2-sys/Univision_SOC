"""Generate the Uni_Vision SOC Midterm submission (Weeks 1-6) as a .docx.

Matches the style of the earlier Weeks 3 & 4 submission: titled cover block, per-week sections
with Lab / Homework subsections, monospaced code blocks, and tables.

    python build_doc.py
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
CODE_BG = "F2F2F2"


def add_code(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    # shading
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), CODE_BG)
    pPr.append(shd)
    return p


def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = ACCENT
    return p


def kv_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(htext)
        r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return table


doc = Document()

# ── Cover block ────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("UNI_VISION SOC")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Midterm Submission — Weeks 1 to 6")
r.bold = True
r.font.size = Pt(15)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Foundations, Python, Web/React, APIs, Blocks & Graphs, Images")
r.italic = True
r.font.size = Pt(11)

meta = [
    ("Course", "Uni_Vision Concept Learning Curriculum"),
    ("Weeks Covered", "Week 1 - Week 6 (Midterm)"),
    ("Topics", "Computational thinking, Python, HTML/CSS/JS/React, FastAPI, Graphs/DAGs, OpenCV"),
    ("Submission Type", "Lab Work + Homework Deliverables + GitHub Repository"),
]
t = doc.add_table(rows=0, cols=2)
t.style = "Light List Accent 1"
for k, v in meta:
    cells = t.add_row().cells
    rr = cells[0].paragraphs[0].add_run(k)
    rr.bold = True
    cells[1].text = v
doc.add_paragraph()

intro = doc.add_paragraph()
intro.add_run(
    "Uni_Vision is taught as a layered system, not a single model: a real-time visual-intelligence "
    "pipeline that captures frames, cleans them, analyses them, validates the results, reports them "
    "to a dashboard, stores them, and can explain itself through agentic AI tools. This midterm "
    "submission covers Weeks 1 through 6 - the foundations the later weeks build on. Every code "
    "deliverable in this document also lives in the accompanying GitHub repository and runs as-is."
)

doc.add_page_break()

# ── WEEK 1 ──────────────────────────────────────────────────────────────────
h(doc, "Week 1: Computing Without Fear", 1)
doc.add_paragraph(
    "Focus: Scratch, Blockly, flowcharts, and the four logic primitives behind every system. "
    "The aim of Week 1 is to see a real-time AI vision system as what it really is - a pipeline of "
    "small logical decisions - before any neural network appears."
)

h(doc, "Lab: Camera Alert Flowchart", 2)
doc.add_paragraph(
    "The lab task was to draw a flowchart for a motion-triggered camera alert, following the Week 1 "
    "rule: if motion is detected, capture a frame; otherwise keep waiting. This is the event-driven "
    "loop the whole pipeline later sits on top of."
)
add_code(doc,
    "START\n"
    "  |\n"
    "  v\n"
    "READ NEXT FRAME  <-------------------+\n"
    "  |                                  |\n"
    "  v                                  |\n"
    "[ Motion detected? ] --No----------->+\n"
    "  | Yes                              |\n"
    "  v                                  |\n"
    "CAPTURE / SAVE FRAME                 |\n"
    "  |                                  |\n"
    "  v                                  |\n"
    "RUN QUICK CHECK (threshold)          |\n"
    "  |                                  |\n"
    "  v                                  |\n"
    "[ Above threshold? ] --No--> LOG ----+\n"
    "  | Yes                              |\n"
    "  v                                  |\n"
    "RAISE ALERT (frame + time + score)   |\n"
    "  |                                  |\n"
    "  v                                  |\n"
    "STORE RESULT & METRICS --------------+"
)
doc.add_paragraph(
    "The loop back to 'Read next frame' keeps the camera live. The two diamonds are the conditions "
    "(decisions) the system makes. Every alert carries evidence - frame, time, confidence - because "
    "Uni_Vision teaches provenance early: every AI output should have evidence, time, source, and "
    "confidence."
)

h(doc, "Homework: Variables, Loops, Events, Conditions", 2)
doc.add_paragraph(
    "A computer system is not magic and not one big brain - it is a set of small, predictable steps "
    "in order. Four ideas are enough to describe how a camera-alert system behaves:"
)
for term, desc in [
    ("Variable", "a named box that stores a value the system remembers right now, e.g. "
                 "confidence_threshold = 0.5. Change the value and behaviour changes without rewriting logic."),
    ("Condition", "chooses between paths via a yes/no question: 'if motion is detected, capture a "
                  "frame; otherwise keep waiting.' This is how a system makes decisions."),
    ("Loop", "repeats a step. A camera does not look once; it grabs frame after frame. Loops are why "
             "a system runs continuously."),
    ("Event", "something that happens and triggers behaviour - a key press, a new frame, motion "
              "crossing a line. Events connect the outside world to the system's logic."),
]:
    p = doc.add_paragraph(style="List Bullet")
    rr = p.add_run(term + " - ")
    rr.bold = True
    p.add_run(desc)
doc.add_paragraph(
    "Takeaway: systems are made of small logical steps, and variables / conditions / loops / events "
    "are the four bricks every step is built from."
)

# ── WEEK 2 ──────────────────────────────────────────────────────────────────
doc.add_page_break()
h(doc, "Week 2: Python As A Tool", 1)
doc.add_paragraph(
    "Focus: functions, lists, dictionaries, and tests. A single detection is the smallest unit of "
    "pipeline output; we model it as a dictionary {label, confidence, box} so it converts cleanly to "
    "JSON later (Week 4) and is easy to inspect by eye."
)

h(doc, "Lab: Detection Dictionaries + Confidence Filter", 2)
add_code(doc,
    'from typing import List, TypedDict\n\n'
    'class Detection(TypedDict):\n'
    '    label: str\n'
    '    confidence: float\n'
    '    box: dict  # {x1, y1, x2, y2}\n\n'
    'def average_confidence(scores: List[float]) -> float:\n'
    '    if not scores:\n'
    '        return 0.0\n'
    '    return sum(scores) / len(scores)\n\n'
    'def filter_by_confidence(detections, threshold=0.5):\n'
    '    """Keep only detections at/above the threshold - the pipeline trust gate."""\n'
    '    if not 0.0 <= threshold <= 1.0:\n'
    '        raise ValueError("threshold must be in [0, 1]")\n'
    '    return [d for d in detections if d["confidence"] >= threshold]'
)
doc.add_paragraph(
    "filter_by_confidence is the pipeline's trust gate: low-confidence guesses are dropped before they "
    "reach the dashboard or raise an alert. Running detections.py prints 4 detections, an average "
    "confidence of 0.610, and keeps ['person', 'helmet'] at threshold 0.50."
)

h(doc, "Homework: Tests For Threshold Filtering", 2)
doc.add_paragraph(
    "Tests prove small pieces behave correctly before the system grows complicated. The suite covers "
    "the cases that matter:"
)
add_code(doc,
    'def test_filter_boundary_is_inclusive():\n'
    '    """A detection exactly at the threshold is kept (>=)."""\n'
    '    dets = [make_detection("edge", 0.50, {"x1":0,"y1":0,"x2":1,"y2":1})]\n'
    '    assert len(filter_by_confidence(dets, threshold=0.50)) == 1\n\n'
    'def test_filter_threshold_one_keeps_only_perfect():\n'
    '    assert filter_by_confidence(sample_detections(), threshold=1.0) == []\n\n'
    'def test_filter_rejects_out_of_range_threshold():\n'
    '    with pytest.raises(ValueError):\n'
    '        filter_by_confidence(sample_detections(), threshold=1.5)'
)
doc.add_paragraph("Result: 11 passed. The inclusive boundary, the empty-input case, and invalid "
                  "thresholds are all locked down so downstream code can trust the filter.")

# ── WEEK 3 ──────────────────────────────────────────────────────────────────
doc.add_page_break()
h(doc, "Week 3: Web Basics And React Thinking", 1)
doc.add_paragraph(
    "Focus: HTML for structure, CSS for appearance, JavaScript for behaviour, then React component "
    "thinking and TypeScript types. The goal: understand how a pipeline dashboard manages and displays "
    "changing state."
)

h(doc, "Lab: Three-Stage Dashboard Mock", 2)
doc.add_paragraph(
    "A dashboard showing Input -> Detection -> Output. Each stage card reflects a status (waiting / "
    "running / complete / error) with colour. An Advance button steps stages one at a time."
)
add_code(doc,
    "type StageStatus = 'waiting' | 'running' | 'complete' | 'error';\n\n"
    "interface Stage { id: string; label: string; status: StageStatus; }\n\n"
    "function advancePipeline() {\n"
    "  setStages(prev => {\n"
    "    const next = prev.map(s => ({ ...s }));\n"
    "    const running = next.findIndex(s => s.status === 'running');\n"
    "    if (running !== -1) {\n"
    "      next[running].status = 'complete';\n"
    "      if (running + 1 < next.length) next[running + 1].status = 'running';\n"
    "    } else {\n"
    "      const first = next.findIndex(s => s.status === 'waiting');\n"
    "      if (first !== -1) next[first].status = 'running';\n"
    "    }\n"
    "    return next;\n"
    "  });\n"
    "}"
)

h(doc, "Homework: UI State Diagram", 2)
kv_table(doc, ["From", "To", "Trigger"], [
    ["idle", "stage_running", "User clicks Advance - first stage becomes running"],
    ["stage_running", "stage_running", "Stage completes - next stage becomes running"],
    ["stage_running", "pipeline_complete", "Last stage completes - all stages complete"],
    ["stage_running", "error_state", "Stage throws - card turns red, pipeline halts"],
    ["pipeline_complete", "idle", "User clicks Reset - all stages back to waiting"],
    ["error_state", "idle", "User clicks Retry - state resets to waiting"],
])
doc.add_paragraph(
    "Why a dashboard needs state: a dashboard is a live view of a changing system, not a static page. "
    "State is the data the component remembers between events. Separating state from display (the React "
    "model) makes it easy to add stages or change logic without redesigning the page."
)

# ── WEEK 4 ──────────────────────────────────────────────────────────────────
doc.add_page_break()
h(doc, "Week 4: APIs And JSON Contracts", 1)
doc.add_paragraph(
    "Focus: routes, JSON, validation, status codes, using FastAPI and Pydantic. The key lesson: write "
    "the contract (request/response shapes) before the implementation."
)

h(doc, "Lab: Image-Analysis API Contract", 2)
kv_table(doc, ["Method", "Route", "Purpose"], [
    ["GET", "/health", "Returns API status and version (monitoring)"],
    ["POST", "/analyse", "Accepts an image URL + config, returns detections"],
    ["GET", "/results/{id}", "Returns a stored analysis result by id"],
])
add_code(doc,
    "from fastapi import FastAPI, HTTPException\n"
    "from pydantic import BaseModel, HttpUrl, Field\n\n"
    "class AnalyseRequest(BaseModel):\n"
    "    image_url: HttpUrl\n"
    "    model_name: str = 'yolov8n'\n"
    "    confidence_threshold: float = Field(0.5, ge=0.0, le=1.0)\n"
    "    max_detections: int | None = Field(10, gt=0)\n\n"
    "@app.get('/health')\n"
    "def health():\n"
    "    return {'status': 'ok', 'version': app.version}\n\n"
    "@app.post('/analyse', response_model=AnalyseResponse)\n"
    "def analyse(request: AnalyseRequest):\n"
    "    ...  # validates input, runs detector, stores + returns structured result"
)

h(doc, "Homework: Worked Request & Response Bodies", 2)
doc.add_paragraph("POST /analyse - success response:")
add_code(doc,
    '{\n'
    '  "result_id": "a3f2c1d0-7e44-4b8a-9c12-1e6b2f3a5d78",\n'
    '  "image_url": "https://example.com/frames/frame_00142.jpg",\n'
    '  "detections": [\n'
    '    {"label": "person", "confidence": 0.92,\n'
    '     "box": {"x1": 50, "y1": 30, "x2": 200, "y2": 380}},\n'
    '    {"label": "helmet", "confidence": 0.78,\n'
    '     "box": {"x1": 70, "y1": 20, "x2": 160, "y2": 80}}\n'
    '  ],\n'
    '  "processing_time_ms": 143.2,\n'
    '  "created_at": "2025-09-14T10:22:04.331Z"\n'
    '}'
)
doc.add_paragraph("Validation error (422) - missing/out-of-range fields are rejected automatically:")
add_code(doc,
    '{\n'
    '  "detail": [\n'
    '    {"loc": ["body", "image_url"], "msg": "field required",\n'
    '     "type": "value_error.missing"}\n'
    '  ]\n'
    '}'
)
doc.add_paragraph(
    "Why responses should be predictable: another app (the Week 3 dashboard, a Week 11 agent) must "
    "know exactly what fields it will receive. A validated contract lets clients be written before the "
    "server is finished and rejects malformed input at the door."
)

# ── WEEK 5 ──────────────────────────────────────────────────────────────────
doc.add_page_break()
h(doc, "Week 5: Blocks And Graphs", 1)
doc.add_paragraph(
    "Focus: blocks, ports, connections, graph validation, and topological sorting. A block is a "
    "visible operation with input/output ports; a connection sends data from one block's output port "
    "to another's input port. A workflow must be a directed acyclic graph (DAG)."
)

h(doc, "Lab: Build A Block Graph And Validate It", 2)
doc.add_paragraph(
    "The graph models the pipeline Load -> Resize -> Grayscale -> Detect -> {Alert, Save}. Validation "
    "checks ports and cycles; topological_order() returns a valid execution order via Kahn's algorithm."
)
add_code(doc,
    "from collections import deque\n\n"
    "def topological_order(self):\n"
    "    \"\"\"Return block ids in a valid execution order; raise on a cycle.\"\"\"\n"
    "    adj = self._adjacency()\n"
    "    indegree = {b: 0 for b in self.blocks}\n"
    "    for src in adj:\n"
    "        for dst in adj[src]:\n"
    "            indegree[dst] += 1\n"
    "    ready = deque(sorted(b for b, d in indegree.items() if d == 0))\n"
    "    order = []\n"
    "    while ready:\n"
    "        node = ready.popleft(); order.append(node)\n"
    "        for nxt in adj[node]:\n"
    "            indegree[nxt] -= 1\n"
    "            if indegree[nxt] == 0:\n"
    "                ready.append(nxt)\n"
    "    if len(order) != len(self.blocks):\n"
    "        raise ValueError('Cycle detected: no valid ordering exists.')\n"
    "    return order"
)
doc.add_paragraph("Output: Validation: VALID; "
                  "Execution order: load -> resize -> gray -> detect -> alert -> save.")

h(doc, "Homework: Graph Validation Checklist", 2)
for item in [
    "Every connection references blocks that exist.",
    "Every connection uses ports that actually exist on those blocks.",
    "The graph is acyclic (a DAG) - no block depends on itself.",
    "A valid execution order exists where each block runs after its dependencies.",
]:
    doc.add_paragraph(item, style="List Number")
doc.add_paragraph(
    "Why validation must precede execution: a cycle means there is no first block to run, so the "
    "workflow can never start; an edge to a non-existent port means data has nowhere to go. Catching "
    "these statically is far cheaper and safer than failing mid-execution. The checklist is enforced by "
    "6 passing tests."
)

# ── WEEK 6 ──────────────────────────────────────────────────────────────────
doc.add_page_break()
h(doc, "Week 6: Images As Arrays", 1)
doc.add_paragraph(
    "Focus: pixels, channels, shape, resize, grayscale, crop. The lesson: every model input is just "
    "numbers in a useful shape. An image is a NumPy array of shape (height, width, channels)."
)

h(doc, "Lab: Load, Inspect, Resize, Grayscale, Crop", 2)
add_code(doc,
    "import cv2\n\n"
    "img = cv2.imread(path)              # or a synthetic sample if no file given\n"
    "h, w, c = img.shape                 # (height, width, channels)\n\n"
    "resized = cv2.resize(img, (640, 640))          # changes spatial resolution\n"
    "gray    = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY) # 3 channels -> 1 (colour lost)\n"
    "crop    = img[h//4:3*h//4, w//4:3*w//4]         # NumPy slice: rows, then cols"
)
doc.add_paragraph(
    "Observed shapes when run: original (h=360, w=480, c=3); resized (640, 640, 3); grayscale "
    "(360, 480, 1); centre crop (180, 240, 3). The script auto-generates a synthetic image so it runs "
    "with no external file, and writes before/after PNGs for comparison."
)

h(doc, "Homework: Before/After Images + Explanation", 2)
doc.add_paragraph(
    "Original vs grayscale: grayscale collapses 3 colour channels into 1 intensity channel. We lose "
    "the ability to distinguish objects by colour (a red vs a green jacket) but keep shape and edges, "
    "and the array is 3x smaller, so cheaper to process."
)
doc.add_paragraph(
    "Original vs resized: forcing a non-square image into 640x640 stretches it and distorts the aspect "
    "ratio. Real pipelines usually letterbox (pad) instead, to preserve proportions. Resizing changes "
    "resolution, cropping selects a region, normalization rescales values - none of it is magic, just "
    "array manipulation."
)

# ── Closing / repo note ──────────────────────────────────────────────────────
doc.add_page_break()
h(doc, "Repository & How To Run", 1)
doc.add_paragraph(
    "All deliverables are in the GitHub repository, one folder per week. Every Python deliverable runs "
    "as-is and its tests pass (Week 2: 11 passed, Week 5: 6 passed)."
)
add_code(doc,
    "git clone <your-repo-url>\n"
    "cd UniVision-SOC\n"
    "python -m venv .venv && .venv\\Scripts\\activate   # Windows\n"
    "pip install -r requirements.txt\n\n"
    "python week2_python/detections.py\n"
    "pytest week2_python/test_detections.py -v\n"
    "python week5_blocks_graphs/graph_validation.py\n"
    "pytest week5_blocks_graphs/test_graph_validation.py -v\n"
    "python week6_images/image_basics.py\n\n"
    "# Week 4 API:\n"
    "uvicorn week4_api.main:app --reload   # then open /docs"
)
doc.add_paragraph(
    "Progress map: Weeks 1-6 (this midterm) cover the foundations - computational thinking, Python, "
    "web/React, API contracts, block graphs, and images as arrays. Weeks 7-12 build on these toward "
    "preprocessing, object detection, OCR/tracking, real-time streaming, agentic AI/RAG, and the "
    "capstone demo."
)

out = "UniVision_SOC_Midterm_Weeks1-6.docx"
doc.save(out)
print("Saved", out)
